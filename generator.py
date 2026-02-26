from __future__ import annotations

import re
import unicodedata
from pathlib import Path
from datetime import datetime
from typing import List, Optional, Any, Dict
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Pt

from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    ListFlowable,
    ListItem,
    Flowable,
)

from ats_sanitize import sanitize_text

OUTPUT_DIR = Path("outputs")


def ensure_outdir() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ============================================================
# OPTIONAL: Intelligence Layer (safe import)
# - generator MUST WORK even if resume_intel.py doesn't exist
# ============================================================
_INTEL_AVAILABLE = False
apply_intelligence = None
enforce_quality = None

try:
    from resume_intel import apply_intelligence as _ai_apply  # type: ignore
    apply_intelligence = _ai_apply
    _INTEL_AVAILABLE = True
except Exception:
    apply_intelligence = None
    _INTEL_AVAILABLE = False

try:
    from resume_intel import enforce_quality as _ai_quality  # type: ignore
    enforce_quality = _ai_quality
except Exception:
    enforce_quality = None


def _apply_intel_safe(master: dict, role_family: str, jd_text: str, jd_struct: Optional[dict]) -> dict:
    """
    Safe wrapper for apply_intelligence with signature flexibility.
    If anything fails -> return master unchanged.
    """
    if not _INTEL_AVAILABLE or apply_intelligence is None:
        return master

    try:
        # Preferred signature (new):
        return apply_intelligence(master, role_family, jd_text, jd_struct)  # type: ignore
    except TypeError:
        # Back-compat: apply_intelligence(master, role_family, jd_text)
        try:
            return apply_intelligence(master, role_family, jd_text)  # type: ignore
        except Exception:
            return master
    except Exception:
        return master


# ----------------------------
# Helpers: wrap long tokens (PDF + DOCX safety)
# ----------------------------
_LONG_TOKEN = re.compile(r"\b[^\s]{28,}\b")


def _soft_break_long_tokens(text: str) -> str:
    if not text:
        return ""
    text = str(text)

    def repl(m):
        tok = m.group(0)
        chunks = [tok[i:i + 14] for i in range(0, len(tok), 14)]
        return "\u200b".join(chunks)

    return _LONG_TOKEN.sub(repl, text)


# ----------------------------
# PDF-safe text cleanup (prevents black squares)
# ----------------------------
_BAD_UNICODE = [
    "\u00ad",
    "\u200b",
    "\u200c",
    "\u200d",
    "\ufeff",
    "\u2060",
]


def pdf_safe_text(s: str) -> str:
    if s is None:
        return ""
    s = str(s)

    for ch in _BAD_UNICODE:
        s = s.replace(ch, "")

    s = (s
         .replace("•", "- ")
         .replace("·", "- ")
         .replace("–", "-")
         .replace("—", "-")
         .replace("−", "-")
         .replace("’", "'")
         .replace("‘", "'")
         .replace("“", '"')
         .replace("”", '"')
         .replace("…", "...")
         .replace("\xa0", " ")
         )

    s = unicodedata.normalize("NFKD", s)
    s = s.encode("ascii", "ignore").decode("ascii")

    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\s+\n", "\n", s)
    s = re.sub(r"\n{3,}", "\n\n", s)

    return s.strip()


# ----------------------------
# Sections ordering (FINAL ORDER)
# ----------------------------
DEFAULT_SECTIONS = [
    "SUMMARY",
    "EDUCATION",
    "SKILLS",
    "EXPERIENCE",
    "PROJECTS",
    "CERTIFICATIONS",
    "PUBLICATIONS",
]


def _normalize_sections(template_sections: Optional[List[str]]) -> List[str]:
    if not template_sections:
        return DEFAULT_SECTIONS
    out: List[str] = []
    seen = set()
    for s in template_sections:
        if not s:
            continue
        k = str(s).strip().upper()
        if k in seen:
            continue
        seen.add(k)
        out.append(k)

    for must in DEFAULT_SECTIONS:
        if must not in out:
            out.append(must)
    return out


# ----------------------------
# Skills formatting (PREVIOUS STYLE = category lines)
# ----------------------------
def _skills_lines(skills: Any) -> List[str]:
    """
    Produces lines like:
      Programming: Python, Java, C++
      Tools: Git, Docker
    (This is the "previous one was better" style.)
    """
    lines: List[str] = []

    if isinstance(skills, dict):
        for cat, items in skills.items():
            cat_name = str(cat).strip() or "Skills"
            if isinstance(items, list):
                clean = []
                seen = set()
                for it in items:
                    s = str(it).strip()
                    if not s:
                        continue
                    k = s.lower()
                    if k in seen:
                        continue
                    seen.add(k)
                    clean.append(s)
                if clean:
                    lines.append(f"{cat_name}: " + ", ".join(clean))
            else:
                s = str(items).strip()
                if s:
                    lines.append(f"{cat_name}: {s}")

    elif isinstance(skills, list):
        clean = []
        seen = set()
        for it in skills:
            s = str(it).strip()
            if not s:
                continue
            k = s.lower()
            if k in seen:
                continue
            seen.add(k)
            clean.append(s)
        if clean:
            lines.append("Skills: " + ", ".join(clean))

    return lines


# ----------------------------
# Make text blocks (shared)
# ----------------------------
def make_resume_text_blocks(
    master: dict,
    role_family: str,
    bridge_info: dict,
    jd_text: str,
    template_sections: Optional[List[str]] = None,
    selected_projects: Optional[List[dict]] = None,
    jd_struct: Optional[dict] = None,
) -> dict:
    # ✅ apply intelligence safely (never crash)
    master2 = _apply_intel_safe(master or {}, role_family, jd_text, jd_struct)

    # ✅ self-learning: persist JD terms into learn_jd store
    try:
        from learn_jd import learn_terms  # type: ignore
        jd_terms = master2.get("_intel_jd_terms") or []
        if isinstance(jd_terms, list) and jd_terms:
            learn_terms([str(x) for x in jd_terms], role_family=role_family)
    except Exception:
        pass


    basics = master2.get("basics", {}) if isinstance(master2.get("basics", {}), dict) else {}
    name = sanitize_text(master2.get("name") or basics.get("name") or "Candidate")
    location = sanitize_text(master2.get("location") or basics.get("location") or "")
    email = sanitize_text(master2.get("email") or basics.get("email") or "")
    phone = sanitize_text(master2.get("phone") or basics.get("phone") or "")
    linkedin = sanitize_text(master2.get("linkedin") or basics.get("linkedin") or "")

    summary = master2.get("summary") or master2.get("objective") or ""
    summary = sanitize_text(summary)

    # Skills lines (previous style)
    skills_obj = master2.get("skills", {}) or {}
    skills_lines = [sanitize_text(x) for x in _skills_lines(skills_obj) if str(x).strip()]

    # Experience/Projects/Education/etc
    experience = master2.get("experience", []) or []
    projects = selected_projects if selected_projects is not None else (master2.get("projects", []) or [])
    education = master2.get("education", []) or []
    publications = master2.get("publications", []) or []
    certifications = master2.get("certifications", []) or master2.get("certs", []) or []

    sections = _normalize_sections(template_sections)

    # Auto-hide empty sections
    def _has_nonempty(x: Any) -> bool:
        if isinstance(x, list):
            return any(str(i).strip() for i in x)
        if isinstance(x, dict):
            return any(_has_nonempty(v) for v in x.values())
        return bool(str(x).strip())

    filtered_sections = []
    for s in sections:
        if s == "PUBLICATIONS" and not _has_nonempty(publications):
            continue
        if s == "CERTIFICATIONS" and not _has_nonempty(certifications):
            continue
        filtered_sections.append(s)

    blocks = {
        "basics": {"name": name, "location": location, "email": email, "phone": phone, "linkedin": linkedin},
        "summary": summary,
        "skills_lines": skills_lines,
        "experience": experience,
        "projects": projects,
        "education": education,
        "certifications": certifications,
        "publications": publications,
        "sections": filtered_sections,
    }

    # Optional quality enforcement if available (safe)
    if enforce_quality is not None:
        try:
            blocks = enforce_quality(blocks, min_words=420)  # type: ignore
        except Exception:
            pass

    return blocks


# ----------------------------
# DOCX generation
# ----------------------------
def generate_docx(
    master: dict,
    role_family: str,
    bridge_info: dict,
    jd_text: str,
    out_name: str,
    template_sections: Optional[List[str]] = None,
    selected_projects: Optional[List[dict]] = None,
    jd_struct: Optional[dict] = None,
) -> Path:
    ensure_outdir()
    out_path = OUTPUT_DIR / f"{out_name}.docx"
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    blocks = make_resume_text_blocks(
        master, role_family, bridge_info, jd_text,
        template_sections=template_sections,
        selected_projects=selected_projects,
        jd_struct=jd_struct
    )

    def h(t: str):
        p = doc.add_paragraph()
        r = p.add_run(t)
        r.bold = True
        r.font.size = Pt(12)

    def bullet(text: str):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(sanitize_text(_soft_break_long_tokens(text)))

    # Header
    header = doc.add_paragraph()
    r = header.add_run(blocks["basics"]["name"])
    r.bold = True
    r.font.size = Pt(16)

    contact = " | ".join([x for x in [
        blocks["basics"]["location"],
        blocks["basics"]["email"],
        blocks["basics"]["phone"],
        blocks["basics"]["linkedin"],
    ] if x])
    if contact:
        doc.add_paragraph(contact)

    doc.add_paragraph("")

    for section in blocks["sections"]:
        if section == "SUMMARY":
            h("Summary")
            doc.add_paragraph(_soft_break_long_tokens(blocks["summary"]))
            doc.add_paragraph("")

        elif section == "EDUCATION":
            if blocks["education"]:
                h("Education")
                for ed in blocks["education"]:
                    if not isinstance(ed, dict):
                        doc.add_paragraph(_soft_break_long_tokens(str(ed)))
                        continue
                    degree = sanitize_text(ed.get("degree", ""))
                    school = sanitize_text(ed.get("school", ""))
                    dates = sanitize_text(ed.get("dates", ""))
                    line = " — ".join([x for x in [degree, school] if x])
                    if dates:
                        line = f"{line} ({dates})"
                    doc.add_paragraph(_soft_break_long_tokens(line))
                doc.add_paragraph("")

        elif section == "SKILLS":
            h("Skills")
            for line in blocks.get("skills_lines") or []:
                doc.add_paragraph(_soft_break_long_tokens(line))
            doc.add_paragraph("")

        elif section == "EXPERIENCE":
            if blocks["experience"]:
                h("Experience")
                for e in blocks["experience"]:
                    if not isinstance(e, dict):
                        doc.add_paragraph(_soft_break_long_tokens(str(e)))
                        continue
                    title = sanitize_text(e.get("title", ""))
                    org = sanitize_text(e.get("org", e.get("company", "")))
                    start = sanitize_text(e.get("start", ""))
                    end = sanitize_text(e.get("end", ""))
                    head = sanitize_text(" — ".join([x for x in [title, org] if x]))
                    dates = sanitize_text(" to ".join([x for x in [start, end] if x]))
                    line = f"{head} ({dates})" if dates else head
                    p = doc.add_paragraph(_soft_break_long_tokens(line))
                    if p.runs:
                        p.runs[0].bold = True
                    for b in e.get("bullets", []) or []:
                        bullet(b)
                doc.add_paragraph("")

        elif section == "PROJECTS":
            if blocks["projects"]:
                h("Projects")
                for pr in blocks["projects"]:
                    if not isinstance(pr, dict):
                        doc.add_paragraph(_soft_break_long_tokens(str(pr)))
                        continue
                    pname = sanitize_text(pr.get("name", "Project"))
                    org = sanitize_text(pr.get("org", ""))
                    tech = sanitize_text(pr.get("tech", ""))          # ✅ ADDED
                    dates = sanitize_text(pr.get("dates", ""))

                    head = " — ".join([x for x in [pname, org] if x])

                    # ✅ show tech stack (don’t hide the intelligence work)
                    if tech:
                        head = f"{head} | Tech: {tech}"

                    if dates:
                        head = f"{head} ({dates})"

                    p = doc.add_paragraph(_soft_break_long_tokens(head))

                    if p.runs:
                        p.runs[0].bold = True
                    for b in pr.get("bullets", []) or []:
                        bullet(b)
                doc.add_paragraph("")

        elif section == "CERTIFICATIONS":
            certs = blocks.get("certifications") or []
            if certs:
                h("Certifications")
                for c in certs:
                    bullet(str(c))
                doc.add_paragraph("")

        elif section == "PUBLICATIONS":
            pubs = blocks.get("publications") or []
            if pubs:
                h("Publications")
                for pub in pubs:
                    if isinstance(pub, dict):
                        title = sanitize_text(pub.get("title", ""))
                        venue = sanitize_text(pub.get("venue", ""))
                        dt = sanitize_text(pub.get("date", ""))
                        line = " — ".join([x for x in [title, venue] if x])
                        if dt:
                            line = f"{line} ({dt})"
                        doc.add_paragraph(_soft_break_long_tokens(line))
                    else:
                        doc.add_paragraph(_soft_break_long_tokens(str(pub)))
                doc.add_paragraph("")

    doc.save(str(out_path))
    return out_path


# ----------------------------
# PDF: divider flowable
# ----------------------------
class HR(Flowable):
    def __init__(self, width=1.0, thickness=1, color=colors.HexColor("#9aa3ad"), space=8):
        super().__init__()
        self.width = width
        self.thickness = thickness
        self.color = color
        self.space = space

    def wrap(self, availWidth, availHeight):
        self.availWidth = availWidth
        return availWidth, self.space

    def draw(self):
        self.canv.saveState()
        self.canv.setStrokeColor(self.color)
        self.canv.setLineWidth(self.thickness)
        y = self.space / 2.0
        self.canv.line(0, y, self.availWidth * self.width, y)
        self.canv.restoreState()


# ----------------------------
# PDF generation (tight margins + 1-page friendly)
# ----------------------------
def generate_pdf_full(
    master: dict,
    role_family: str,
    bridge_info: dict,
    jd_text: str,
    out_name: str,
    template_sections: Optional[List[str]] = None,
    selected_projects: Optional[List[dict]] = None,
    jd_struct: Optional[dict] = None,
) -> Path:
    ensure_outdir()
    out_path = OUTPUT_DIR / f"{out_name}.pdf"

    blocks = make_resume_text_blocks(
        master, role_family, bridge_info, jd_text,
        template_sections=template_sections,
        selected_projects=selected_projects,
        jd_struct=jd_struct
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "Title",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=15,
        leading=16,
        spaceAfter=4,
    )
    header_style = ParagraphStyle(
        "Header",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=11,
        textColor=colors.black,
        spaceAfter=6,
    )
    h_style = ParagraphStyle(
        "Section",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=12,
        spaceBefore=6,
        spaceAfter=4,
    )
    p_style = ParagraphStyle(
        "P",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=11,
        spaceAfter=2,
    )

    # ✅ tight margins
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=letter,
        rightMargin=0.45 * inch,
        leftMargin=0.45 * inch,
        topMargin=0.45 * inch,
        bottomMargin=0.45 * inch,
        title=pdf_safe_text(blocks["basics"]["name"]),
        author=pdf_safe_text(blocks["basics"]["name"]),
    )

    story: List[Any] = []

    def safe_para(txt: str, style=p_style):
        txt = sanitize_text(txt)
        txt = _soft_break_long_tokens(txt)
        txt = pdf_safe_text(txt)
        txt = escape(txt)
        return Paragraph(txt, style)

    def bullets(items: List[str], left_indent: int = 14):
        clean: List[str] = []
        for x in items:
            s = sanitize_text(str(x))
            s = _soft_break_long_tokens(s)
            s = pdf_safe_text(s)
            s = escape(s)
            if s.strip():
                clean.append(s)
        if not clean:
            return
        lf = ListFlowable(
            [ListItem(Paragraph(x, p_style), leftIndent=left_indent) for x in clean],
            bulletType="bullet",
            leftIndent=left_indent,
            bulletFontName="Helvetica",
            bulletFontSize=8.5,
        )
        story.append(lf)

    def divider():
        story.append(Spacer(1, 2))
        story.append(HR(width=1.0, thickness=1, color=colors.HexColor("#b8c0c8"), space=8))
        story.append(Spacer(1, 1))

    # Header
    story.append(safe_para(blocks["basics"]["name"], title_style))
    contact = " | ".join([x for x in [
        blocks["basics"]["location"],
        blocks["basics"]["email"],
        blocks["basics"]["phone"],
        blocks["basics"]["linkedin"],
    ] if x])
    if contact:
        story.append(safe_para(contact, header_style))

    for sec in blocks["sections"]:
        if sec == "SUMMARY":
            story.append(safe_para("Summary", h_style))
            story.append(safe_para(blocks["summary"]))
            divider()

        elif sec == "EDUCATION":
            if blocks["education"]:
                story.append(safe_para("Education", h_style))
                for ed in blocks["education"]:
                    if isinstance(ed, dict):
                        degree = sanitize_text(ed.get("degree", ""))
                        school = sanitize_text(ed.get("school", ""))
                        dates = sanitize_text(ed.get("dates", ""))
                        line = " — ".join([x for x in [degree, school] if x])
                        if dates:
                            line = f"{line} ({dates})"
                        story.append(safe_para(line))
                    else:
                        story.append(safe_para(str(ed)))
                divider()

        elif sec == "SKILLS":
            story.append(safe_para("Skills", h_style))
            for line in blocks.get("skills_lines") or []:
                story.append(safe_para(line))
            divider()

        elif sec == "EXPERIENCE":
            if blocks["experience"]:
                story.append(safe_para("Experience", h_style))
                for e in blocks["experience"]:
                    if not isinstance(e, dict):
                        story.append(safe_para(str(e)))
                        continue
                    title = sanitize_text(e.get("title", ""))
                    org = sanitize_text(e.get("org", e.get("company", "")))
                    start = sanitize_text(e.get("start", ""))
                    end = sanitize_text(e.get("end", ""))
                    head = " — ".join([x for x in [title, org] if x])
                    dates = " to ".join([x for x in [start, end] if x])
                    if dates:
                        head = f"{head} ({dates})"
                    story.append(safe_para(head, ParagraphStyle("JobHead", parent=p_style, fontName="Helvetica-Bold")))
                    bullets(e.get("bullets", []) or [], left_indent=14)
                    story.append(Spacer(1, 3))
                divider()

        elif sec == "PROJECTS":
            if blocks["projects"]:
                story.append(safe_para("Projects", h_style))
                for pr in blocks["projects"]:
                    if not isinstance(pr, dict):
                        story.append(safe_para(str(pr)))
                        continue
                    pname = sanitize_text(pr.get("name", "Project"))
                    org = sanitize_text(pr.get("org", ""))
                    tech = sanitize_text(pr.get("tech", ""))          # ✅ ADDED
                    dates = sanitize_text(pr.get("dates", ""))

                    head = " — ".join([x for x in [pname, org] if x])

                    # ✅ show tech stack
                    if tech:
                        head = f"{head} | Tech: {tech}"

                    if dates:
                        head = f"{head} ({dates})"

                    story.append(safe_para(head, ParagraphStyle("ProjHead", parent=p_style, fontName="Helvetica-Bold")))

                    bullets(pr.get("bullets", []) or [], left_indent=14)
                    story.append(Spacer(1, 3))
                divider()

        elif sec == "CERTIFICATIONS":
            certs = blocks.get("certifications") or []
            if certs:
                story.append(safe_para("Certifications", h_style))
                bullets([str(c) for c in certs], left_indent=14)
                divider()

        elif sec == "PUBLICATIONS":
            pubs = blocks.get("publications") or []
            if pubs:
                story.append(safe_para("Publications", h_style))
                for pub in pubs:
                    if isinstance(pub, dict):
                        title = sanitize_text(pub.get("title", ""))
                        venue = sanitize_text(pub.get("venue", ""))
                        dt = sanitize_text(pub.get("date", ""))
                        line = " — ".join([x for x in [title, venue] if x])
                        if dt:
                            line = f"{line} ({dt})"
                        story.append(safe_para(line))
                    else:
                        story.append(safe_para(str(pub)))
                divider()

    doc.build(story)
    return out_path


# ----------------------------
# Cover letter
# ----------------------------
def generate_cover_letter(master: dict, company: str, role_title: str, out_name: str, *args, **kwargs) -> Path:
    ensure_outdir()
    out_path = OUTPUT_DIR / f"{out_name}_cover_letter.docx"
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    basics = master.get("basics", {}) if isinstance(master.get("basics", {}), dict) else {}
    name = sanitize_text(master.get("name") or basics.get("name") or "Candidate")

    doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    doc.add_paragraph("")
    doc.add_paragraph(sanitize_text(f"Re: {role_title} at {company}"))
    doc.add_paragraph("")

    body = sanitize_text(
        f"Dear Hiring Manager,\n\n"
        f"I’m applying for the {role_title} role at {company}. I bring hands-on project experience in software development, "
        f"strong fundamentals in Java/Python/SQL, and collaborative delivery.\n\n"
        f"I’d welcome the chance to discuss how I can contribute.\n\n"
        f"Sincerely,\n{name}"
    )

    for para in body.split("\n\n"):
        doc.add_paragraph(_soft_break_long_tokens(para))

    doc.save(str(out_path))
    return out_path


# Keep recruiter_message import compatibility
def recruiter_message(master: dict, company: str, role_title: str) -> str:
    basics = master.get("basics", {}) if isinstance(master.get("basics", {}), dict) else {}
    name = master.get("name") or basics.get("name") or "Candidate"
    return sanitize_text(
        f"Hi {company} team — I’m {name}. I applied for the {role_title} role and would love to connect. "
        f"I can contribute with strong foundations in software development and project delivery. "
        f"Open to a quick chat if helpful."
    )
